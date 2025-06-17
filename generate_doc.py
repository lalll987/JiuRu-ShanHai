from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_code_document():
    # 创建新的Word文档
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('研究论文写作助手项目代码文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加项目说明
    doc.add_paragraph('本文档包含研究论文写作助手项目的所有源代码。项目由以下模块组成：')
    modules = [
        '1. 文献综述模块 (literature_review_module)',
        '2. 研究模块 (research_module)',
        '3. 写作模块 (writing_module)',
        '4. 评估模块 (evaluation_module)'
    ]
    for module in modules:
        doc.add_paragraph(module, style='List Bullet')
    
    # 遍历项目目录
    for root, dirs, files in os.walk('.'):
        # 跳过venv和.git目录
        if 'venv' in root or '.git' in root:
            continue
            
        # 处理Python文件
        for file in files:
            if file.endswith('.py'):
                file_path = os.path.join(root, file)
                relative_path = os.path.relpath(file_path, '.')
                
                # 添加文件标题
                doc.add_heading(f'文件：{relative_path}', level=1)
                
                # 读取文件内容
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # 添加代码内容
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
                # 添加分隔线
                doc.add_paragraph('=' * 80)
    
    # 保存文档
    doc.save('project_code_documentation.docx')

if __name__ == '__main__':
    create_code_document() 